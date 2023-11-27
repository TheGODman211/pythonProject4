
import openpyxl
from docx import Document



# Load the workbook
wb = openpyxl.load_workbook("D:/excel/dde/no effective cut non life- cocobod_dollar bonds.xlsx")

# Get all sheet names
sheet_names = wb.sheetnames
names = []
list = []
# Print sheet names
# for sheet_name in sheet_names:
#     names.append(sheet_name)
#
# for i in names:
#     wb.active = wb[i]
#     list.append(wb.active.cell(row=1, column=1).value)
#
#
# names = []

# Assuming you have already loaded the workbook into the wb variable
for sheet in wb:
    if sheet.sheet_state != 'hidden':  # Only consider visible sheets
        names.append(sheet.title)

# List to hold cell values
values = []

for name in names:
    sheet = wb[name]
    values.append(sheet.cell(row=1, column=1).value)

# Now, values list contains the values from cell A1 from each visible sheet
print(values)

print(names)
print(list)




# Create a new Word document
# doc = Document()
#
# # Add a title
# doc.add_heading('Sheet Names and Values', level=1)
#
# # Loop through names and values to add them to the document
# for name, value in zip(names, values):
#     # Add sheet name
#     doc.add_paragraph(f"Sheet Name: {name}")
#
#     # Add value
#     doc.add_paragraph(f"Value: {value}")
#
#     # Add a line break for better readability
#     doc.add_paragraph()
#
# # Save the document
# doc.save('D:/non life-sheet_names_and_values.docx')


# Create a new Excel workbook and select the active sheet
wb = openpyxl.Workbook()
sheet = wb.active

# Set column headers (optional)
sheet['A1'] = "Sheet Name"
sheet['B1'] = "Value"

# Write data to the Excel sheet
for idx, (names, values) in enumerate(zip(names, values), start=2):
    sheet[f"A{idx}"] = names
    # sheet[f"B{idx}"] = values

# Save the workbook
wb.save("D:/excel/NONlifelist.xlsx")





